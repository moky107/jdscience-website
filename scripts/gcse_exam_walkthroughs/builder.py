"""Build GCSE exam walkthrough PDF and DOCX packs."""

from __future__ import annotations

from pathlib import Path
from typing import Any

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Mm, Pt, RGBColor
from reportlab.lib.units import mm
from reportlab.platypus import (
    HRFlowable,
    KeepTogether,
    PageBreak,
    Paragraph,
    SimpleDocTemplate,
    Spacer,
)

from .command_words import COMMAND_WORDS, intro_paragraphs
from .theme import OUT, draw_branded_page, draw_cover_page, page_margins, styles


def _bullets(st, items: list[str], style="bullet"):
    return [Paragraph(f"• {item}", st[style]) for item in items]


def _section(st, title: str):
    return [Paragraph(title, st["section"]), Spacer(1, 2 * mm)]


def _h2(st, title: str):
    return [Paragraph(title, st["h2"]), Spacer(1, 1 * mm)]


def _h3(st, title: str):
    return [Paragraph(title, st["h3"])]


def _body(st, text: str):
    return [Paragraph(text, st["body"])]


def _numbered(st, items: list[str]):
    out = []
    for i, item in enumerate(items, 1):
        out.append(Paragraph(f"{i}. {item}", st["bullet"]))
    return out


def build_command_word_section(st) -> list:
    flow = _section(st, "Section 1: Command word guide")
    for cw in COMMAND_WORDS:
        flow.extend(_h2(st, cw["word"]))
        flow.append(Paragraph(f"<b>What you must do:</b> {cw['must_do']}", st["body"]))
        flow.append(Paragraph(f"<b>Common mistake:</b> {cw['mistake']}", st["body"]))
        flow.append(Paragraph(f"<b>Model sentence starter:</b> {cw['starter']}", st["body"]))
        flow.append(Spacer(1, 2 * mm))
    flow.append(PageBreak())
    return flow


def build_intro(st, subject: str) -> list:
    flow = _section(st, "Introduction")
    for para in intro_paragraphs(subject):
        flow.append(Paragraph(para, st["body"]))
        flow.append(Spacer(1, 2 * mm))
    flow.append(PageBreak())
    return flow


def build_question_walkthrough(st, q: dict[str, Any], number: int) -> list:
    flow = []
    meta = (
        f"<b>Topic:</b> {q['topic']} &nbsp;|&nbsp; "
        f"<b>Command word:</b> {q['command']} &nbsp;|&nbsp; "
        f"<b>Marks:</b> {q['marks']} &nbsp;|&nbsp; "
        f"<b>Skill tested:</b> {q['skill']}"
    )
    block = [
        Paragraph(f"QUESTION {number}", st["h2"]),
        Paragraph(q["text"], st["question"]),
        Paragraph("QUESTION TYPE", st["label"]),
        Paragraph(meta, st["meta"]),
        Paragraph("WHAT THE QUESTION IS ASKING", st["label"]),
        Paragraph(q["asking"], st["body"]),
        Paragraph("STEP-BY-STEP THINKING", st["label"]),
    ]
    block.extend(_numbered(st, q["steps"]))
    block.append(Paragraph("MODEL ANSWER", st["label"]))
    block.append(Paragraph(q["model"], st["body"]))
    block.append(Paragraph("MARK BREAKDOWN", st["label"]))
    for i, mark in enumerate(q["marks_breakdown"], 1):
        block.append(Paragraph(f"Mark {i}: {mark}", st["bullet"]))
    block.append(Paragraph("COMMON MISTAKES", st["label"]))
    block.extend(_bullets(st, q["mistakes"]))
    block.append(Paragraph("EXAMINER TIP", st["label"]))
    block.append(Paragraph(q["tip"], st["body"]))
    if q.get("extension"):
        block.append(Paragraph("GRADE 7–9 EXTENSION", st["label"]))
        block.append(Paragraph(q["extension"], st["body"]))
    block.append(Spacer(1, 4 * mm))
    block.append(HRFlowable(width="100%", thickness=0.5, color=styles()["body"].textColor))
    block.append(Spacer(1, 3 * mm))
    flow.append(KeepTogether(block[:8]))
    flow.extend(block[8:])
    return flow


def build_topic_section(st, topic: str, questions: list[dict], start_num: int) -> tuple[list, int]:
    flow = _section(st, f"Section 2: {topic}")
    num = start_num
    for q in questions:
        flow.extend(build_question_walkthrough(st, q, num))
        num += 1
    flow.append(PageBreak())
    return flow, num


def build_calculation_section(st, calcs: list[dict]) -> list:
    flow = _section(st, "Section 3: Calculation walkthroughs")
    for i, calc in enumerate(calcs, 1):
        flow.extend(_h2(st, f"Calculation {i}: {calc['title']}"))
        flow.append(Paragraph(calc.get("question", ""), st["body"]))
        for label, key in [
            ("Formula", "formula"),
            ("Rearrangement", "rearrange"),
            ("Substitution", "substitution"),
            ("Working", "working"),
            ("Answer", "answer"),
        ]:
            if calc.get(key):
                flow.append(Paragraph(f"<b>{label}:</b> {calc[key]}", st["body"]))
        if calc.get("note"):
            flow.append(Paragraph(calc["note"], st["note"]))
        flow.append(Spacer(1, 3 * mm))
    flow.append(PageBreak())
    return flow


def build_practical_section(st, practicals: list[dict]) -> list:
    flow = _section(st, "Section 4: Required practical walkthroughs")
    for i, prac in enumerate(practicals, 1):
        flow.extend(_h2(st, f"Practical {i}: {prac['title']}"))
        for label, key in [
            ("Aim", "aim"),
            ("Variables", "variables"),
            ("Method", "method"),
            ("Risk assessment", "risk"),
            ("Results table", "results"),
            ("Graph guidance", "graph"),
            ("Conclusion", "conclusion"),
            ("Evaluation", "evaluation"),
        ]:
            val = prac.get(key)
            if val:
                flow.append(Paragraph(f"<b>{label}:</b> {val}", st["body"]))
        if prac.get("exam_questions"):
            flow.append(Paragraph("<b>Common exam questions:</b>", st["label"]))
            flow.extend(_bullets(st, prac["exam_questions"]))
        if prac.get("model_answers"):
            flow.append(Paragraph("<b>Model answers:</b>", st["label"]))
            flow.extend(_bullets(st, prac["model_answers"]))
        flow.append(Spacer(1, 4 * mm))
    flow.append(PageBreak())
    return flow


def build_mini_mock(st, mock: dict) -> list:
    flow = _section(st, "Section 5: Full mini mock paper")
    flow.append(Paragraph(f"<b>Total marks: {mock['total_marks']}</b>", st["body"]))
    flow.append(Spacer(1, 2 * mm))
    for i, q in enumerate(mock["questions"], 1):
        marks = q.get("marks", "?")
        flow.append(Paragraph(f"<b>Question {i}</b> ({marks} marks)", st["h3"]))
        flow.append(Paragraph(q["text"], st["body"]))
        flow.append(Spacer(1, 2 * mm))
    flow.append(PageBreak())
    flow.append(Paragraph("Mini mock — step-by-step walkthrough answers", st["section"]))
    for i, ans in enumerate(mock["answers"], 1):
        flow.extend(_h2(st, f"Answer {i}"))
        flow.append(Paragraph(ans.get("walkthrough", ans.get("model", "")), st["body"]))
        if ans.get("marks_breakdown"):
            flow.append(Paragraph("<b>Mark breakdown:</b>", st["label"]))
            for j, m in enumerate(ans["marks_breakdown"], 1):
                flow.append(Paragraph(f"Mark {j}: {m}", st["bullet"]))
        flow.append(Spacer(1, 2 * mm))
    return flow


def build_pdf(pack: dict, out_path: Path) -> Path:
    st = styles()
    subject = pack["subject"]
    title = f"GCSE {subject} Exam Walkthrough Pack"
    out_path.parent.mkdir(parents=True, exist_ok=True)

    doc = SimpleDocTemplate(
        str(out_path),
        pagesize=(210 * mm, 297 * mm),
        title=title,
        author="JDScience",
        **page_margins(),
    )

    story: list = []
    story.append(Spacer(1, 180 * mm))
    story.append(PageBreak())
    story.extend(build_intro(st, subject))
    story.extend(build_command_word_section(st))

    qnum = 1
    for topic_block in pack["topics"]:
        block, qnum = build_topic_section(st, topic_block["name"], topic_block["questions"], qnum)
        story.extend(block)

    story.extend(build_calculation_section(st, pack["calculations"]))
    story.extend(build_practical_section(st, pack["practicals"]))
    story.extend(build_mini_mock(st, pack["mini_mock"]))

    def first_page(c, d):
        draw_cover_page(c, d, subject)

    def later_pages(c, d):
        draw_branded_page(c, d, title)

    doc.build(story, onFirstPage=first_page, onLaterPages=later_pages)
    return out_path


def _docx_heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    return p


def build_docx(pack: dict, out_path: Path) -> Path:
    subject = pack["subject"]
    title = f"GCSE {subject} Exam Walkthrough Pack"
    doc = Document()
    section = doc.sections[0]
    section.page_height = Mm(297)
    section.page_width = Mm(210)
    section.top_margin = Mm(18)
    section.bottom_margin = Mm(16)
    section.left_margin = Mm(16)
    section.right_margin = Mm(16)

    cover = doc.add_paragraph()
    cover.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = cover.add_run(f"GCSE {subject}\nExam Walkthrough Pack")
    run.bold = True
    run.font.size = Pt(24)
    run.font.color.rgb = RGBColor(0, 77, 64)
    sub = doc.add_paragraph(
        "Step-by-step worked exam-style questions with full explanations\n"
        "Original JDScience exam-style questions — not copied from past papers\n"
        "JDScience · jdscience.co.uk · Produced by JDScience"
    )
    sub.runs[0].font.size = Pt(11)
    doc.add_page_break()

    _docx_heading(doc, "Introduction", 1)
    for para in intro_paragraphs(subject):
        doc.add_paragraph(para)

    _docx_heading(doc, "Section 1: Command word guide", 1)
    for cw in COMMAND_WORDS:
        _docx_heading(doc, cw["word"], 2)
        doc.add_paragraph(f"What you must do: {cw['must_do']}")
        doc.add_paragraph(f"Common mistake: {cw['mistake']}")
        doc.add_paragraph(f"Model sentence starter: {cw['starter']}")

    qnum = 1
    for topic_block in pack["topics"]:
        _docx_heading(doc, f"Section 2: {topic_block['name']}", 1)
        for q in topic_block["questions"]:
            _docx_heading(doc, f"Question {qnum}", 2)
            doc.add_paragraph(q["text"])
            doc.add_paragraph(
                f"Topic: {q['topic']} | Command: {q['command']} | Marks: {q['marks']} | Skill: {q['skill']}"
            )
            doc.add_paragraph("What the question is asking").runs[0].bold = True
            doc.add_paragraph(q["asking"])
            doc.add_paragraph("Step-by-step thinking").runs[0].bold = True
            for i, step in enumerate(q["steps"], 1):
                doc.add_paragraph(f"{i}. {step}", style="List Number")
            doc.add_paragraph("Model answer").runs[0].bold = True
            doc.add_paragraph(q["model"])
            doc.add_paragraph("Mark breakdown").runs[0].bold = True
            for i, m in enumerate(q["marks_breakdown"], 1):
                doc.add_paragraph(f"Mark {i}: {m}")
            doc.add_paragraph("Common mistakes").runs[0].bold = True
            for m in q["mistakes"]:
                doc.add_paragraph(m, style="List Bullet")
            doc.add_paragraph("Examiner tip").runs[0].bold = True
            doc.add_paragraph(q["tip"])
            if q.get("extension"):
                doc.add_paragraph("Grade 7–9 extension").runs[0].bold = True
                doc.add_paragraph(q["extension"])
            qnum += 1

    _docx_heading(doc, "Section 3: Calculation walkthroughs", 1)
    for i, calc in enumerate(pack["calculations"], 1):
        _docx_heading(doc, f"Calculation {i}: {calc['title']}", 2)
        for key in ("question", "formula", "rearrange", "substitution", "working", "answer", "note"):
            if calc.get(key):
                doc.add_paragraph(f"{key.capitalize()}: {calc[key]}")

    _docx_heading(doc, "Section 4: Required practical walkthroughs", 1)
    for prac in pack["practicals"]:
        _docx_heading(doc, prac["title"], 2)
        for key in ("aim", "variables", "method", "risk", "results", "graph", "conclusion", "evaluation"):
            if prac.get(key):
                doc.add_paragraph(f"{key.replace('_', ' ').title()}: {prac[key]}")

    _docx_heading(doc, "Section 5: Full mini mock paper", 1)
    mock = pack["mini_mock"]
    doc.add_paragraph(f"Total marks: {mock['total_marks']}")
    for i, q in enumerate(mock["questions"], 1):
        doc.add_paragraph(f"Question {i} ({q.get('marks', '?')} marks): {q['text']}")
    _docx_heading(doc, "Mini mock walkthrough answers", 2)
    for i, ans in enumerate(mock["answers"], 1):
        doc.add_paragraph(f"Answer {i}: {ans.get('walkthrough', ans.get('model', ''))}")

    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out_path))
    return out_path
